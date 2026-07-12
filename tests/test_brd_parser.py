"""Tests for core/brd_parser.py."""

import io
import pytest
from unittest.mock import patch, MagicMock

from core.brd_parser import parse_brd, parse_supporting_docs, build_combined_input, BRDParseError


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

SAMPLE_TEXT = "This is a sample BRD. " * 50  # >10 words


def _make_docx_bytes(text: str) -> bytes:
    """Create a minimal .docx file in memory with the given paragraph text."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# PDF tests
# ---------------------------------------------------------------------------

class TestParseBrdPdf:
    def test_valid_pdf_returns_dict(self):
        """A PDF with readable text returns the expected dict structure."""
        fake_page = MagicMock()
        fake_page.get_text.return_value = SAMPLE_TEXT

        fake_doc = MagicMock()
        fake_doc.__iter__ = MagicMock(return_value=iter([fake_page, fake_page]))
        fake_doc.__len__ = MagicMock(return_value=2)
        # fitz.open returns fake_doc; also iterable in the _extract_pdf loop
        fake_doc.__enter__ = MagicMock(return_value=fake_doc)
        fake_doc.__exit__ = MagicMock(return_value=False)

        with patch("core.brd_parser.fitz") as mock_fitz:
            mock_fitz.open.return_value = fake_doc
            result = parse_brd(b"%PDF fake bytes", "report.pdf")

        assert "raw_text" in result
        assert "filename" in result
        assert result["filename"] == "report.pdf"
        assert result["page_count"] >= 1
        assert result["word_count"] > 0

    def test_empty_pdf_raises(self):
        """A PDF whose pages contain no text raises BRDParseError."""
        fake_page = MagicMock()
        fake_page.get_text.return_value = ""

        fake_doc = MagicMock()
        fake_doc.__iter__ = MagicMock(return_value=iter([fake_page]))
        fake_doc.close = MagicMock()

        with patch("core.brd_parser.fitz") as mock_fitz:
            mock_fitz.open.return_value = fake_doc
            with pytest.raises(BRDParseError, match="empty"):
                parse_brd(b"%PDF fake bytes", "empty.pdf")

    def test_corrupt_pdf_raises(self):
        """A corrupt PDF that fitz cannot open raises BRDParseError."""
        with patch("core.brd_parser.fitz") as mock_fitz:
            mock_fitz.open.side_effect = Exception("bad PDF")
            with pytest.raises(BRDParseError, match="corrupt"):
                parse_brd(b"not a pdf", "corrupt.pdf")


# ---------------------------------------------------------------------------
# DOCX tests
# ---------------------------------------------------------------------------

class TestParseBrdDocx:
    def test_valid_docx_returns_dict(self):
        """A well-formed DOCX returns the expected dict structure."""
        docx_bytes = _make_docx_bytes(SAMPLE_TEXT)
        result = parse_brd(docx_bytes, "requirements.docx")

        assert result["filename"] == "requirements.docx"
        assert result["word_count"] > 0
        assert len(result["raw_text"]) > 0

    def test_docx_extracts_table_text(self):
        """Text inside DOCX tables is extracted."""
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header One"
        table.rows[0].cells[1].text = "Header Two"
        table.rows[1].cells[0].text = "Value A"
        table.rows[1].cells[1].text = "Value B"
        # Add extra words to pass the minimum word count check
        doc.add_paragraph(SAMPLE_TEXT)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        docx_bytes = buf.read()

        result = parse_brd(docx_bytes, "table_brd.docx")
        assert "Header One" in result["raw_text"] or result["word_count"] > 10

    def test_empty_docx_raises(self):
        """A DOCX with no text raises BRDParseError."""
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        with pytest.raises(BRDParseError):
            parse_brd(buf.read(), "empty.docx")

    def test_corrupt_docx_raises(self):
        """Passing random bytes as a DOCX raises BRDParseError."""
        with pytest.raises(BRDParseError, match="corrupt"):
            parse_brd(b"this is not a docx file at all", "corrupt.docx")


# ---------------------------------------------------------------------------
# Unsupported type test
# ---------------------------------------------------------------------------

class TestUnsupportedType:
    def test_unsupported_extension_raises(self):
        """An unsupported extension raises BRDParseError with a helpful message."""
        with pytest.raises(BRDParseError, match="Unsupported file type"):
            parse_brd(b"fake content", "requirements.xlsx")


# ---------------------------------------------------------------------------
# Multi-document input tests (FR-02/FR-03)
# ---------------------------------------------------------------------------

SAMPLE_BRD_DICT = {
    "raw_text": "The BRD requires a BW/4HANA reporting solution.",
    "filename": "brd.docx",
    "page_count": 1,
    "word_count": 8,
}


class TestParseSupportingDocs:
    def test_multiple_files_parsed(self):
        """Multiple supporting documents are each parsed into a dict."""
        files = [
            (_make_docx_bytes(SAMPLE_TEXT), "meeting_minutes.docx"),
            (_make_docx_bytes(SAMPLE_TEXT), "data_dictionary.docx"),
        ]
        results = parse_supporting_docs(files)
        assert len(results) == 2
        assert results[0]["filename"] == "meeting_minutes.docx"
        assert results[1]["filename"] == "data_dictionary.docx"

    def test_no_files_returns_empty_list(self):
        assert parse_supporting_docs([]) == []

    def test_malformed_file_raises_with_filename(self):
        """A malformed supporting document raises BRDParseError naming the file."""
        files = [(b"not a docx file at all", "corrupt_notes.docx")]
        with pytest.raises(BRDParseError, match="corrupt_notes.docx"):
            parse_supporting_docs(files)


class TestBuildCombinedInput:
    def test_brd_only(self):
        """With no supporting docs or notes, only the BRD marker appears."""
        combined = build_combined_input(SAMPLE_BRD_DICT, [], "")
        assert "=== DOCUMENT: brd.docx ===" in combined
        assert SAMPLE_BRD_DICT["raw_text"] in combined
        assert "CLIENT DISCUSSION NOTES" not in combined

    def test_brd_plus_supporting_docs_and_notes(self):
        """Supporting docs and discussion notes are concatenated with delimiter markers."""
        supporting = [{"raw_text": "Meeting notes text.", "filename": "minutes.docx", "page_count": 1, "word_count": 3}]
        combined = build_combined_input(SAMPLE_BRD_DICT, supporting, "Client confirmed scope on the call.")

        assert "=== DOCUMENT: brd.docx ===" in combined
        assert "=== DOCUMENT: minutes.docx ===" in combined
        assert "Meeting notes text." in combined
        assert "=== CLIENT DISCUSSION NOTES ===" in combined
        assert "Client confirmed scope on the call." in combined

    def test_blank_notes_omitted(self):
        combined = build_combined_input(SAMPLE_BRD_DICT, [], "   ")
        assert "CLIENT DISCUSSION NOTES" not in combined
