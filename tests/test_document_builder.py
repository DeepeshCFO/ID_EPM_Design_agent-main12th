"""Tests for core/document_builder.py."""

import io
import re

from docx import Document

from core.document_builder import assemble_document, parse_batch_sections


SAMPLE_METADATA = {"consultant": "Jane Doe", "client": "Acme Corp", "project": "EPM Rollout", "engagement_code": "ENG-001"}

SIMPLE_SECTIONS = [
    {"number": "1", "title": "Document Control", "description": "", "format": "table"},
    {"number": "2", "title": "Executive Summary", "description": "", "format": "prose"},
    {"number": "3", "title": "Reporting Catalogue", "description": "", "format": "table"},
]

TABLE_CONTENT = (
    "| Report | Type | Frequency |\n"
    "| --- | --- | --- |\n"
    "| P&L Report | Operational | Daily |\n"
    "| Forecast | Planning | Monthly |\n"
    "| Variance | Management | Weekly |\n"
    "| Consolidation | Statutory | Monthly |\n"
)

SECTION_CONTENT = {
    "1": "| Item | Detail |\n| --- | --- |\n| Version | 1.0 |\n",
    "2": "This is the executive summary body text describing the project.",
    "3": TABLE_CONTENT,
}


def _assemble():
    docx_bytes = assemble_document(SECTION_CONTENT, SIMPLE_SECTIONS, doc_type="FSD", technology="SAP BW/4HANA", metadata=SAMPLE_METADATA)
    return Document(io.BytesIO(docx_bytes)), docx_bytes


def _cell_fill(cell) -> str:
    match = re.search(r'w:fill="([0-9A-Fa-f]{6})"', cell._tc.xml)
    return match.group(1).upper() if match else ""


# ---------------------------------------------------------------------------
# parse_batch_sections
# ---------------------------------------------------------------------------

class TestParseBatchSections:
    def test_splits_on_section_markers(self):
        response = "## SECTION_1:\nContent one.\n## SECTION_2:\nContent two."
        result = parse_batch_sections(response)
        assert result["1"] == "Content one."
        assert result["2"] == "Content two."

    def test_no_markers_returns_empty_dict(self):
        assert parse_batch_sections("No markers here.") == {}


# ---------------------------------------------------------------------------
# Named styles
# ---------------------------------------------------------------------------

class TestNamedStyles:
    def test_all_epm_styles_registered(self):
        doc, _ = _assemble()
        style_names = {s.name for s in doc.styles}
        expected = {
            "EPM Heading 1", "EPM Heading 2", "EPM Heading 3", "EPM Body Text",
            "EPM Table Header", "EPM Table Body", "EPM Code Block", "EPM Bullet", "EPM Caption",
        }
        assert expected.issubset(style_names)

    def test_section_headings_use_epm_heading_1(self):
        doc, _ = _assemble()
        heading_paragraphs = [p for p in doc.paragraphs if p.text.startswith("2. Executive Summary")]
        assert heading_paragraphs
        assert heading_paragraphs[0].style.name == "EPM Heading 1"


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

class TestCoverPage:
    def test_cover_page_contains_required_fields(self):
        doc, _ = _assemble()
        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert SAMPLE_METADATA["client"] in all_text
        assert SAMPLE_METADATA["project"] in all_text
        assert "Functional Specification Document" in all_text
        assert "SAP BW/4HANA" in all_text
        assert SAMPLE_METADATA["consultant"] in all_text
        assert SAMPLE_METADATA["engagement_code"] in all_text
        assert "Confidential" in all_text


# ---------------------------------------------------------------------------
# Table of Contents field
# ---------------------------------------------------------------------------

class TestTableOfContents:
    def test_toc_field_present_in_document_xml(self):
        doc, _ = _assemble()
        xml = doc.element.xml
        assert "TOC" in xml
        assert "Table of Contents" in "\n".join(p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# Table shading — header + alternating rows
# ---------------------------------------------------------------------------

class TestTableShading:
    def test_header_row_shaded_sap_blue(self):
        doc, _ = _assemble()
        table = doc.tables[-1]  # the reporting catalogue table (last one built)
        for cell in table.rows[0].cells:
            assert _cell_fill(cell) == "1F497D"

    def test_alternating_body_rows_shaded(self):
        doc, _ = _assemble()
        table = doc.tables[-1]
        # row 1 (first data row) -> white, row 2 -> light blue, row 3 -> white, row 4 -> light blue
        assert _cell_fill(table.rows[1].cells[0]) == "FFFFFF"
        assert _cell_fill(table.rows[2].cells[0]) == "D6E4F0"
        assert _cell_fill(table.rows[3].cells[0]) == "FFFFFF"
        assert _cell_fill(table.rows[4].cells[0]) == "D6E4F0"


# ---------------------------------------------------------------------------
# Missing section content
# ---------------------------------------------------------------------------

class TestMissingSectionContent:
    def test_missing_section_renders_placeholder(self):
        docx_bytes = assemble_document({"1": "Only section 1."}, SIMPLE_SECTIONS, doc_type="TSD", technology="SAP BW/4HANA", metadata=SAMPLE_METADATA)
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[Section content not generated]" in all_text
