"""Tests for core/tsd_generator.py."""

import io
import pytest
from unittest.mock import patch, MagicMock

from core.tsd_generator import (
    resolve_section_structure,
    split_into_batches,
    generate_tsd_batch,
    build_tsd_bytes,
)
from templates.tsd_default_structure import TSD_DEFAULT_SECTIONS


SAMPLE_STRUCTURED_SUMMARY = {
    "requirements": ["Daily P&L reporting on BW/4HANA"],
    "kpis": ["Revenue", "EBITDA"],
    "data_sources": ["S/4HANA"],
    "domain": ["finance"],
    "key_entities": ["Legal Entity"],
    "open_gaps": [],
}

SAMPLE_SAP_SKILL = {
    "summary": "SAP BW/4HANA",
    "key_concepts": ["aDSO"],
    "kpis": ["Revenue"],
    "sap_objects": ["InfoObject"],
    "common_patterns": ["Delta extraction"],
}

SAMPLE_DOMAIN_SKILLS = [
    {"domain": "finance", "summary": "Finance", "key_concepts": [], "kpis": [], "sap_objects": [], "common_patterns": []}
]

SAMPLE_METADATA = {"consultant": "Jane Doe", "client": "Acme Corp", "project": "EPM Rollout", "engagement_code": "ENG-001"}

SAMPLE_FSD_FULL_TEXT = (
    "1. Document Control\nVersion 1.0\n\n5. Solution Architecture\nBW/4HANA aDSO layer feeding SAC stories.\n"
)


def _mock_jinja_env(rendered: str = "rendered prompt"):
    mock_template = MagicMock()
    mock_template.render.return_value = rendered
    mock_env = MagicMock()
    mock_env.get_template.return_value = mock_template
    return mock_env


# ---------------------------------------------------------------------------
# resolve_section_structure
# ---------------------------------------------------------------------------

class TestResolveSectionStructureTsd:
    def test_no_template_returns_default(self):
        sections = resolve_section_structure(None)
        assert sections == TSD_DEFAULT_SECTIONS

    def test_template_with_headings_uses_template(self):
        from docx import Document

        doc = Document()
        doc.add_heading("Technical Architecture", level=1)
        doc.add_heading("Data Modelling", level=1)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        sections = resolve_section_structure(buf.read())
        assert len(sections) == 2
        assert sections[0]["title"] == "Technical Architecture"


# ---------------------------------------------------------------------------
# split_into_batches
# ---------------------------------------------------------------------------

class TestSplitIntoBatchesTsd:
    def test_default_batch_size_splits_14_sections_into_5_batches(self, monkeypatch):
        monkeypatch.delenv("GENERATION_BATCH_SIZE", raising=False)
        batches = split_into_batches(TSD_DEFAULT_SECTIONS)
        assert len(batches) == 5
        assert [len(b) for b in batches] == [3, 3, 3, 3, 2]


# ---------------------------------------------------------------------------
# generate_tsd_batch
# ---------------------------------------------------------------------------

class TestGenerateTsdBatch:
    @patch("core.tsd_generator.call_llm")
    @patch("core.tsd_generator.get_jinja_env")
    def test_returns_llm_response_for_one_batch(self, mock_env, mock_llm):
        mock_env.return_value = _mock_jinja_env()
        mock_llm.return_value = "## SECTION_1:\nContent."

        batches = split_into_batches(TSD_DEFAULT_SECTIONS)
        result = generate_tsd_batch(
            fsd_full_text=SAMPLE_FSD_FULL_TEXT,
            technology="SAP BW/4HANA",
            sap_skill=SAMPLE_SAP_SKILL,
            domain_skills=SAMPLE_DOMAIN_SKILLS,
            structured_summary=SAMPLE_STRUCTURED_SUMMARY,
            clarification_answers={},
            metadata=SAMPLE_METADATA,
            section_structure=TSD_DEFAULT_SECTIONS,
            batch_sections=batches[0],
        )

        assert result == "## SECTION_1:\nContent."
        mock_llm.assert_called_once()

    @patch("core.tsd_generator.call_llm")
    @patch("core.tsd_generator.get_jinja_env")
    def test_fsd_full_text_passed_to_prompt_render(self, mock_env, mock_llm):
        mock_template = MagicMock()
        mock_template.render.return_value = "rendered"
        mock_env.return_value.get_template.return_value = mock_template
        mock_llm.return_value = "## SECTION_1:\nContent."

        batches = split_into_batches(TSD_DEFAULT_SECTIONS)
        generate_tsd_batch(
            fsd_full_text=SAMPLE_FSD_FULL_TEXT,
            technology="SAP BW/4HANA",
            sap_skill=SAMPLE_SAP_SKILL,
            domain_skills=SAMPLE_DOMAIN_SKILLS,
            structured_summary=SAMPLE_STRUCTURED_SUMMARY,
            clarification_answers={},
            metadata=SAMPLE_METADATA,
            section_structure=TSD_DEFAULT_SECTIONS,
            batch_sections=batches[0],
        )

        render_calls = mock_template.render.call_args_list
        tsd_call = next((c for c in render_calls if "fsd_full_text" in c.kwargs), None)
        assert tsd_call is not None
        assert tsd_call.kwargs["fsd_full_text"] == SAMPLE_FSD_FULL_TEXT

    @patch("core.tsd_generator.call_llm")
    @patch("core.tsd_generator.get_jinja_env")
    def test_locked_context_defaults_to_empty_string(self, mock_env, mock_llm):
        """Skip Review fast-mode hand-off (CLAUDE.md §3.7): a cold-start batch call
        (no prior interactive-loop sections locked) must render an empty locked_context,
        never None or a missing kwarg."""
        mock_template = MagicMock()
        mock_template.render.return_value = "rendered"
        mock_env.return_value.get_template.return_value = mock_template
        mock_llm.return_value = "## SECTION_1:\nContent."

        batches = split_into_batches(TSD_DEFAULT_SECTIONS)
        generate_tsd_batch(
            fsd_full_text=SAMPLE_FSD_FULL_TEXT,
            technology="SAP BW/4HANA",
            sap_skill=SAMPLE_SAP_SKILL,
            domain_skills=SAMPLE_DOMAIN_SKILLS,
            structured_summary=SAMPLE_STRUCTURED_SUMMARY,
            clarification_answers={},
            metadata=SAMPLE_METADATA,
            section_structure=TSD_DEFAULT_SECTIONS,
            batch_sections=batches[0],
        )

        render_kwargs = mock_template.render.call_args.kwargs
        assert render_kwargs["locked_context"] == ""

    @patch("core.tsd_generator.call_llm")
    @patch("core.tsd_generator.get_jinja_env")
    def test_locked_context_passed_through_to_prompt_render(self, mock_env, mock_llm):
        mock_template = MagicMock()
        mock_template.render.return_value = "rendered"
        mock_env.return_value.get_template.return_value = mock_template
        mock_llm.return_value = "## SECTION_1:\nContent."

        batches = split_into_batches(TSD_DEFAULT_SECTIONS)
        generate_tsd_batch(
            fsd_full_text=SAMPLE_FSD_FULL_TEXT,
            technology="SAP BW/4HANA",
            sap_skill=SAMPLE_SAP_SKILL,
            domain_skills=SAMPLE_DOMAIN_SKILLS,
            structured_summary=SAMPLE_STRUCTURED_SUMMARY,
            clarification_answers={},
            metadata=SAMPLE_METADATA,
            section_structure=TSD_DEFAULT_SECTIONS,
            batch_sections=batches[0],
            locked_context="### Section 1: Document Control (locked)\nAlready approved content.",
        )

        render_kwargs = mock_template.render.call_args.kwargs
        assert render_kwargs["locked_context"] == "### Section 1: Document Control (locked)\nAlready approved content."


# ---------------------------------------------------------------------------
# build_tsd_bytes
# ---------------------------------------------------------------------------

class TestBuildTsdBytes:
    def test_assembles_valid_docx_bytes(self):
        section_content = {s["number"]: f"Content for {s['title']}." for s in TSD_DEFAULT_SECTIONS}

        result = build_tsd_bytes(section_content, TSD_DEFAULT_SECTIONS, "SAP BW/4HANA", SAMPLE_METADATA)

        assert isinstance(result, bytes)
        assert len(result) > 0
