"""Tests for core/fsd_generator.py."""

import io
import pytest
from unittest.mock import patch, MagicMock

from core.fsd_generator import (
    resolve_section_structure,
    split_into_batches,
    generate_fsd_batch,
    build_fsd_bytes,
)
from templates.fsd_default_structure import FSD_DEFAULT_SECTIONS


SAMPLE_STRUCTURED_SUMMARY = {
    "requirements": ["Daily P&L reporting on BW/4HANA"],
    "kpis": ["Revenue", "EBITDA"],
    "data_sources": ["S/4HANA"],
    "domain": ["finance"],
    "key_entities": ["Legal Entity"],
    "open_gaps": [],
}

SAMPLE_SAP_SKILL = {
    "summary": "SAP BW/4HANA",
    "key_concepts": ["aDSO"],
    "kpis": ["Revenue"],
    "sap_objects": ["InfoObject"],
    "common_patterns": ["Delta extraction"],
}

SAMPLE_DOMAIN_SKILLS = [
    {"domain": "finance", "summary": "Finance", "key_concepts": [], "kpis": [], "sap_objects": [], "common_patterns": []}
]

SAMPLE_METADATA = {"consultant": "Jane Doe", "client": "Acme Corp", "project": "EPM Rollout", "engagement_code": "ENG-001"}


def _mock_jinja_env(rendered: str = "rendered prompt"):
    mock_template = MagicMock()
    mock_template.render.return_value = rendered
    mock_env = MagicMock()
    mock_env.get_template.return_value = mock_template
    return mock_env


# ---------------------------------------------------------------------------
# resolve_section_structure
# ---------------------------------------------------------------------------

class TestResolveSectionStructure:
    def test_no_template_returns_default(self):
        sections = resolve_section_structure(None)
        assert sections == FSD_DEFAULT_SECTIONS

    def test_template_with_valid_headings_uses_template(self):
        from docx import Document

        doc = Document()
        doc.add_heading("Executive Summary", level=1)
        doc.add_heading("Data Requirements", level=1)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        template_bytes = buf.read()

        sections = resolve_section_structure(template_bytes)
        assert len(sections) == 2
        assert sections[0]["title"] == "Executive Summary"

    def test_template_with_no_headings_falls_back_to_default(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("No headings here, just body text.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        sections = resolve_section_structure(buf.read())
        assert sections == FSD_DEFAULT_SECTIONS


# ---------------------------------------------------------------------------
# split_into_batches
# ---------------------------------------------------------------------------

class TestSplitIntoBatches:
    def test_default_batch_size_splits_14_sections_into_5_batches(self, monkeypatch):
        monkeypatch.delenv("GENERATION_BATCH_SIZE", raising=False)
        batches = split_into_batches(FSD_DEFAULT_SECTIONS)
        assert len(batches) == 5
        assert [len(b) for b in batches] == [3, 3, 3, 3, 2]

    def test_custom_batch_size(self, monkeypatch):
        monkeypatch.setenv("GENERATION_BATCH_SIZE", "5")
        batches = split_into_batches(FSD_DEFAULT_SECTIONS)
        assert len(batches) == 3
        assert [len(b) for b in batches] == [5, 5, 4]

    def test_batches_cover_every_section_in_order(self):
        batches = split_into_batches(FSD_DEFAULT_SECTIONS)
        flattened = [s for batch in batches for s in batch]
        assert flattened == FSD_DEFAULT_SECTIONS


# ---------------------------------------------------------------------------
# generate_fsd_batch
# ---------------------------------------------------------------------------

class TestGenerateFsdBatch:
    @patch("core.fsd_generator.call_llm")
    @patch("core.fsd_generator.get_jinja_env")
    def test_returns_llm_response_for_one_batch(self, mock_env, mock_llm):
        mock_env.return_value = _mock_jinja_env()
        mock_llm.return_value = "## SECTION_1:\nContent."

        batches = split_into_batches(FSD_DEFAULT_SECTIONS)
        result = generate_fsd_batch(
            structured_summary=SAMPLE_STRUCTURED_SUMMARY,
            technology="SAP BW/4HANA",
            sap_skill=SAMPLE_SAP_SKILL,
            domain_skills=SAMPLE_DOMAIN_SKILLS,
            clarification_answers={},
            metadata=SAMPLE_METADATA,
            section_structure=FSD_DEFAULT_SECTIONS,
            batch_sections=batches[0],
        )

        assert result == "## SECTION_1:\nContent."
        mock_llm.assert_called_once()

    @patch("core.fsd_generator.call_llm")
    @patch("core.fsd_generator.get_jinja_env")
    def test_uses_max_tokens_batch_env_var(self, mock_env, mock_llm, monkeypatch):
        monkeypatch.setenv("MAX_TOKENS_BATCH", "1234")
        mock_env.return_value = _mock_jinja_env()
        mock_llm.return_value = "## SECTION_1:\nContent."

        batches = split_into_batches(FSD_DEFAULT_SECTIONS)
        generate_fsd_batch(
            structured_summary=SAMPLE_STRUCTURED_SUMMARY,
            technology="SAP BW/4HANA",
            sap_skill=SAMPLE_SAP_SKILL,
            domain_skills=SAMPLE_DOMAIN_SKILLS,
            clarification_answers=None,
            metadata=SAMPLE_METADATA,
            section_structure=FSD_DEFAULT_SECTIONS,
            batch_sections=batches[0],
        )

        assert mock_llm.call_args.kwargs["max_tokens"] == 1234

    @patch("core.fsd_generator.call_llm")
    @patch("core.fsd_generator.get_jinja_env")
    def test_locked_context_defaults_to_empty_string(self, mock_env, mock_llm):
        """Skip Review fast-mode hand-off (CLAUDE.md §3.7): a cold-start batch call
        (no prior interactive-loop sections locked) must render an empty locked_context,
        never None or a missing kwarg."""
        mock_template = MagicMock()
        mock_template.render.return_value = "rendered"
        mock_env.return_value.get_template.return_value = mock_template
        mock_llm.return_value = "## SECTION_1:\nContent."

        batches = split_into_batches(FSD_DEFAULT_SECTIONS)
        generate_fsd_batch(
            structured_summary=SAMPLE_STRUCTURED_SUMMARY,
            technology="SAP BW/4HANA",
            sap_skill=SAMPLE_SAP_SKILL,
            domain_skills=SAMPLE_DOMAIN_SKILLS,
            clarification_answers={},
            metadata=SAMPLE_METADATA,
            section_structure=FSD_DEFAULT_SECTIONS,
            batch_sections=batches[0],
        )

        render_kwargs = mock_template.render.call_args.kwargs
        assert render_kwargs["locked_context"] == ""

    @patch("core.fsd_generator.call_llm")
    @patch("core.fsd_generator.get_jinja_env")
    def test_locked_context_passed_through_to_prompt_render(self, mock_env, mock_llm):
        mock_template = MagicMock()
        mock_template.render.return_value = "rendered"
        mock_env.return_value.get_template.return_value = mock_template
        mock_llm.return_value = "## SECTION_1:\nContent."

        batches = split_into_batches(FSD_DEFAULT_SECTIONS)
        generate_fsd_batch(
            structured_summary=SAMPLE_STRUCTURED_SUMMARY,
            technology="SAP BW/4HANA",
            sap_skill=SAMPLE_SAP_SKILL,
            domain_skills=SAMPLE_DOMAIN_SKILLS,
            clarification_answers={},
            metadata=SAMPLE_METADATA,
            section_structure=FSD_DEFAULT_SECTIONS,
            batch_sections=batches[0],
            locked_context="### Section 1: Document Control (locked)\nAlready approved content.",
        )

        render_kwargs = mock_template.render.call_args.kwargs
        assert render_kwargs["locked_context"] == "### Section 1: Document Control (locked)\nAlready approved content."


# ---------------------------------------------------------------------------
# build_fsd_bytes
# ---------------------------------------------------------------------------

class TestBuildFsdBytes:
    def test_assembles_valid_docx_bytes(self):
        section_content = {s["number"]: f"Content for {s['title']}." for s in FSD_DEFAULT_SECTIONS}

        result = build_fsd_bytes(section_content, FSD_DEFAULT_SECTIONS, "SAP BW/4HANA", SAMPLE_METADATA)

        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_missing_section_content_does_not_raise(self):
        section_content = {"1": "Only section 1 has content."}
        result = build_fsd_bytes(section_content, FSD_DEFAULT_SECTIONS, "SAP BW/4HANA", SAMPLE_METADATA)
        assert isinstance(result, bytes)
