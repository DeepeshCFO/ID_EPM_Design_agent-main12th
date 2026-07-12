"""Word document assembly — all python-docx logic lives here. Named styles only,
no inline font/colour overrides (CLAUDE.md §3.6, FSD §8).
"""

import io
import logging
import re
from datetime import date

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

logger = logging.getLogger(__name__)

_SECTION_MARKER_RE = re.compile(r"##\s*SECTION_(\d+)\s*:", re.IGNORECASE)
_TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|[-| :]+\|\s*$")

_SAP_BLUE = RGBColor(0x1F, 0x49, 0x7D)
_DARK_GREY = RGBColor(0x59, 0x59, 0x59)
_LIGHT_GREY = "D9D9D9"
_ALT_ROW_FILL = "D6E4F0"

_DOC_TYPE_LABELS = {
    "FSD": "Functional Specification Document",
    "TSD": "Technical Specification Document",
}

DEFAULT_METADATA: dict = {
    "consultant": "SAP EPM Design Agent",
    "client": "Client",
    "project": "SAP EPM Implementation",
    "engagement_code": "",
}


class DocumentBuildError(Exception):
    """Raised when document assembly fails."""


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_batch_sections(llm_response: str) -> dict:
    """Split one batch's LLM response at ## SECTION_N: markers. Returns {section_number: content}."""
    sections = {}
    current_num = None
    current_lines = []

    for line in llm_response.splitlines():
        match = _SECTION_MARKER_RE.match(line)
        if match:
            if current_num is not None:
                sections[current_num] = "\n".join(current_lines).strip()
            current_num = match.group(1)
            current_lines = []
        else:
            if current_num is not None:
                current_lines.append(line)

    if current_num is not None:
        sections[current_num] = "\n".join(current_lines).strip()

    return sections


def extract_template_structure(template_bytes: bytes) -> list:
    """Extract section structure from a .docx template's headings.

    Returns a list of dicts matching the FSD/TSD section schema.
    Falls back to an empty list if the template has no recognisable headings.
    """
    try:
        doc = Document(io.BytesIO(template_bytes))
    except Exception as exc:
        logger.warning("Could not read template structure: %s", exc)
        return []

    sections = []
    number = 0
    for para in doc.paragraphs:
        style_name = para.style.name if para.style is not None else ""
        if style_name.startswith("Heading"):
            number += 1
            sections.append({
                "number": str(number),
                "title": para.text.strip(),
                "description": para.text.strip(),
                "format": "prose",
            })

    logger.info("Extracted %d sections from template", len(sections))
    return sections


def assemble_document(
    section_content_dict: dict,
    section_structure: list,
    doc_type: str,
    technology: str,
    metadata: dict | None = None,
) -> bytes:
    """Assemble the full Word document from section content + structure.

    Single entrypoint used for first-pass generation, batch retries, and section
    regeneration alike — CLAUDE.md §3.6: document_builder owns ALL python-docx logic.
    """
    metadata = {**DEFAULT_METADATA, **(metadata or {})}
    doc_label = _DOC_TYPE_LABELS.get(doc_type, doc_type)

    doc = Document()
    _register_epm_styles(doc)
    _add_header_footer(doc, doc_label, metadata)

    _add_cover_page(doc, doc_label, doc_type, technology, metadata)
    doc.add_page_break()

    if section_structure:
        first_section, remaining_sections = section_structure[0], section_structure[1:]
        _render_section(doc, first_section, section_content_dict)
        doc.add_page_break()

        _add_toc_field(doc)
        doc.add_page_break()

        for section in remaining_sections:
            _render_section(doc, section, section_content_dict)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info("Document assembled: %s (%d sections)", doc_type, len(section_structure))
    return buf.read()


# ---------------------------------------------------------------------------
# Cover page, header/footer, TOC
# ---------------------------------------------------------------------------

def _add_cover_page(doc: Document, doc_label: str, doc_type: str, technology: str, metadata: dict) -> None:
    """Add the page-1 cover page (FSD §8.1)."""
    client_para = doc.add_paragraph(style=doc.styles["EPM Heading 1"])
    client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client_para.add_run(metadata["client"])

    project_para = doc.add_paragraph(style=doc.styles["EPM Heading 2"])
    project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project_para.add_run(metadata["project"])

    doc_type_para = doc.add_paragraph(style=doc.styles["EPM Body Text"])
    doc_type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_type_para.add_run(doc_label).bold = True

    tech_para = doc.add_paragraph(style=doc.styles["EPM Body Text"])
    tech_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tech_para.add_run(f"SAP Technology: {technology}")

    version_para = doc.add_paragraph(style=doc.styles["EPM Body Text"])
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_para.add_run("Version 1.0 — Draft")

    date_para = doc.add_paragraph(style=doc.styles["EPM Body Text"])
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Generated: {date.today().strftime('%d %B %Y')}")

    author_para = doc.add_paragraph(style=doc.styles["EPM Body Text"])
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run(f"Author: {metadata['consultant']}")

    if metadata.get("engagement_code"):
        eng_para = doc.add_paragraph(style=doc.styles["EPM Body Text"])
        eng_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eng_para.add_run(f"Engagement Code: {metadata['engagement_code']}")

    conf_para = doc.add_paragraph(style=doc.styles["EPM Caption"])
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_para.add_run("Confidential — Internal Use Only")


def _add_toc_field(doc: Document) -> None:
    """Insert a Word auto-generated Table of Contents field (page 3, FSD §8.3)."""
    heading = doc.add_paragraph("Table of Contents", style=doc.styles["EPM Heading 1"])
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    para = doc.add_paragraph()
    run = para.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and select 'Update Field' to generate the Table of Contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def _add_header_footer(doc: Document, doc_label: str, metadata: dict) -> None:
    """Add page header (title + version + client, right-aligned) and footer
    (confidentiality left, page numbers right) to all sections (FSD §8.3)."""
    for section in doc.sections:
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = ""
        header_para.style = doc.styles["EPM Caption"]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para.add_run(f"{doc_label} — Version 1.0 — {metadata['client']}")

        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = ""
        footer_para.style = doc.styles["EPM Caption"]
        footer_para.paragraph_format.tab_stops.add_tab_stop(
            section.page_width - section.left_margin - section.right_margin,
            WD_TAB_ALIGNMENT.RIGHT,
        )
        footer_para.add_run("Confidentiality — Internal Use Only\t")
        _add_page_number(footer_para)


def _add_page_number(paragraph) -> None:
    """Insert a 'Page X of Y' PAGE/NUMPAGES field run into the given paragraph."""
    run = paragraph.add_run("Page ")

    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)

    run2 = paragraph.add_run(" of ")

    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "begin")
    instr_text2 = OxmlElement("w:instrText")
    instr_text2.text = "NUMPAGES"
    fld_char4 = OxmlElement("w:fldChar")
    fld_char4.set(qn("w:fldCharType"), "end")
    run2._r.append(fld_char3)
    run2._r.append(instr_text2)
    run2._r.append(fld_char4)


# ---------------------------------------------------------------------------
# Section rendering
# ---------------------------------------------------------------------------

def _render_section(doc: Document, section: dict, section_content_dict: dict) -> None:
    """Render one top-level section (heading + body) from the content dict."""
    num = section["number"]
    title = section["title"]
    content = (section_content_dict.get(num) or "").strip()

    heading = doc.add_paragraph(f"{num}. {title}", style=doc.styles["EPM Heading 1"])
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if content:
        _render_section_content(doc, content)
    else:
        logger.warning("No content generated for section %s: %s", num, title)
        placeholder = doc.add_paragraph(style=doc.styles["EPM Body Text"])
        placeholder.add_run("[Section content not generated]").italic = True


def _render_section_content(doc: Document, content: str) -> None:
    """Render section content — tables, bullet lists, code blocks, and prose."""
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        # Markdown table detection — collect consecutive table rows
        if _TABLE_ROW_RE.match(line):
            table_lines = []
            while i < len(lines) and (_TABLE_ROW_RE.match(lines[i]) or _TABLE_SEP_RE.match(lines[i])):
                if not _TABLE_SEP_RE.match(lines[i]):
                    table_lines.append(lines[i])
                i += 1
            if table_lines:
                _create_table_from_markdown(doc, table_lines)
            continue

        # Subheading (## or ###)
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style=doc.styles["EPM Heading 3"])
            i += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style=doc.styles["EPM Heading 2"])
            i += 1
            continue

        # Bullet list
        if line.strip().startswith(("- ", "* ", "+ ")):
            doc.add_paragraph(line.strip()[2:], style=doc.styles["EPM Bullet"])
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", line.strip()):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line.strip()), style=doc.styles["EPM Bullet"])
            i += 1
            continue

        # Non-empty prose line
        if line.strip():
            doc.add_paragraph(line.strip(), style=doc.styles["EPM Body Text"])

        i += 1


def _add_code_block(doc: Document, code: str) -> None:
    """Add a monospaced code block paragraph using the EPM Code Block style."""
    doc.add_paragraph(code, style=doc.styles["EPM Code Block"])


def _create_table_from_markdown(doc: Document, table_lines: list) -> None:
    """Convert markdown pipe-delimited rows into a Word table with named styles."""
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    col_count = max(len(row) for row in rows)
    rows = [row + [""] * (col_count - len(row)) for row in rows]

    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        is_header = r_idx == 0
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            para.style = doc.styles["EPM Table Header"] if is_header else doc.styles["EPM Table Body"]
            para.add_run(cell_text)

    _apply_table_styles(table)


def _apply_table_styles(table) -> None:
    """Apply header shading and alternating row shading to an assembled table (FSD §8.2)."""
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            fill = "1F497D"
        elif r_idx % 2 == 0:
            fill = _ALT_ROW_FILL
        else:
            fill = "FFFFFF"
        for cell in row.cells:
            _shade_cell(cell, fill)


def _shade_cell(cell, fill_hex: str) -> None:
    """Apply a background shade to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


# ---------------------------------------------------------------------------
# Named style registration (CLAUDE.md §3.6 / FSD §8.2)
# ---------------------------------------------------------------------------

def _register_epm_styles(doc: Document) -> None:
    """Register all EPM named styles on the document at creation time."""
    styles = doc.styles

    _add_paragraph_style(
        styles, "EPM Heading 1", base="Heading 1",
        font_name="Arial", size=16, bold=True, color=_SAP_BLUE, space_before=12,
    )
    _add_paragraph_style(
        styles, "EPM Heading 2", base="Heading 2",
        font_name="Arial", size=13, bold=True, color=_SAP_BLUE, space_before=8,
    )
    _add_paragraph_style(
        styles, "EPM Heading 3", base="Heading 3",
        font_name="Arial", size=11, bold=True, color=_DARK_GREY, space_before=6,
    )
    _add_paragraph_style(
        styles, "EPM Body Text", base="Normal",
        font_name="Arial", size=10, space_after=6,
    )
    _add_paragraph_style(
        styles, "EPM Table Header", base="Normal",
        font_name="Arial", size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
    )
    _add_paragraph_style(
        styles, "EPM Table Body", base="Normal",
        font_name="Arial", size=10,
    )
    _add_paragraph_style(
        styles, "EPM Code Block", base="Normal",
        font_name="Courier New", size=9, left_indent=Inches(0.5), shading=_LIGHT_GREY,
    )
    _add_paragraph_style(
        styles, "EPM Bullet", base="List Bullet",
        font_name="Arial", size=10,
    )
    _add_paragraph_style(
        styles, "EPM Caption", base="Normal",
        font_name="Arial", size=9, italic=True, color=_DARK_GREY,
    )


def _add_paragraph_style(
    styles,
    name: str,
    base: str,
    font_name: str,
    size: int,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    space_before: int | None = None,
    space_after: int | None = None,
    left_indent=None,
    shading: str | None = None,
):
    """Add (or replace) one named paragraph style based on a built-in style."""
    if name in styles:
        style = styles[name]
    else:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        if base in styles:
            style.base_style = styles[base]

    style.font.name = font_name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    if color is not None:
        style.font.color.rgb = color
    if space_before is not None:
        style.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        style.paragraph_format.space_after = Pt(space_after)
    if left_indent is not None:
        style.paragraph_format.left_indent = left_indent
    if shading is not None:
        pPr = style.element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), shading)
        pPr.append(shd)

    return style
