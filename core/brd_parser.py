"""BRD document ingestion — PDF and DOCX extraction."""

import io
import logging
from pathlib import Path

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

logger = logging.getLogger(__name__)

_SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
_MIN_WORD_COUNT = 10


class BRDParseError(Exception):
    """Raised when the BRD cannot be parsed or is unusable."""


def parse_brd(file_bytes: bytes, filename: str) -> dict:
    """Extract text from a BRD file and return a structured dict.

    Returns {"raw_text": str, "filename": str, "page_count": int, "word_count": int}.
    Raises BRDParseError for unsupported types, empty files, or corrupt content.
    """
    ext = Path(filename).suffix.lower()
    if ext not in _SUPPORTED_EXTENSIONS:
        raise BRDParseError(
            f"Unsupported file type '{ext}'. Please upload a PDF (.pdf) or Word (.docx) document."
        )

    logger.info("Parsing BRD: %s (ext=%s, bytes=%d)", filename, ext, len(file_bytes))

    if ext == ".pdf":
        raw_text, page_count = _extract_pdf(file_bytes)
    else:
        raw_text, page_count = _extract_docx(file_bytes)

    raw_text = raw_text.strip()
    word_count = len(raw_text.split())

    if word_count < _MIN_WORD_COUNT:
        raise BRDParseError(
            f"The uploaded file appears to be empty or contains too little text (found {word_count} words). "
            "Please upload a complete Business Requirements Document."
        )

    logger.info("BRD parsed: %d pages, %d words", page_count, word_count)

    return {
        "raw_text": raw_text,
        "filename": filename,
        "page_count": page_count,
        "word_count": word_count,
    }


def _extract_pdf(file_bytes: bytes) -> tuple:
    """Extract text from a PDF using PyMuPDF. Returns (text, page_count)."""
    if fitz is None:
        raise BRDParseError("PDF parsing requires PyMuPDF. Run: pip install pymupdf")

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        raise BRDParseError(f"Could not open the PDF file. It may be corrupt or password-protected. ({exc})") from exc

    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()

    return "\n".join(pages), len(pages)


def parse_supporting_docs(files: list) -> list:
    """Parse a list of (file_bytes, filename) supporting documents.

    Returns a list of dicts in the same shape as parse_brd(). Raises BRDParseError
    (with the filename in the message) if any individual document fails to parse.
    """
    parsed = []
    for file_bytes, filename in files:
        try:
            parsed.append(parse_brd(file_bytes, filename))
        except BRDParseError as exc:
            raise BRDParseError(f"Supporting document '{filename}': {exc}") from exc
    return parsed


def build_combined_input(brd_dict: dict, supporting_docs: list, discussion_notes: str) -> str:
    """Concatenate the BRD, supporting documents, and discussion notes into one text block.

    Documents are separated by '=== DOCUMENT: filename ===' markers; discussion notes
    (if present) are appended under a '=== CLIENT DISCUSSION NOTES ===' marker.
    """
    parts = [f"=== DOCUMENT: {brd_dict['filename']} ===\n{brd_dict['raw_text']}"]

    for doc in supporting_docs or []:
        parts.append(f"=== DOCUMENT: {doc['filename']} ===\n{doc['raw_text']}")

    if discussion_notes and discussion_notes.strip():
        parts.append(f"=== CLIENT DISCUSSION NOTES ===\n{discussion_notes.strip()}")

    return "\n\n".join(parts)


def _extract_docx(file_bytes: bytes) -> tuple:
    """Extract text from a DOCX using python-docx. Returns (text, page_count)."""
    try:
        from docx import Document
    except ImportError as exc:
        raise BRDParseError("DOCX parsing requires python-docx. Run: pip install python-docx") from exc

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise BRDParseError(f"Could not open the Word document. It may be corrupt or in an unsupported format. ({exc})") from exc

    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    # Extract text from tables as well
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    # DOCX does not reliably expose page count — use section count as proxy
    page_count = max(1, len(doc.sections))

    return "\n".join(paragraphs), page_count
